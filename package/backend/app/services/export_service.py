import html
import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def _iter_blocks(text: str):
    for raw_block in re.split(r"\n\s*\n", text.strip()):
        block = raw_block.strip()
        if not block:
            continue
        match = re.match(r"^(#{1,6})\s+(.+)$", block, flags=re.S)
        if match:
            yield "heading", len(match.group(1)), match.group(2).strip()
        else:
            yield "body", 0, block


def build_docx(text: str) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(30)
    section.right_margin = Mm(25)

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.space_after = Pt(6)

    for kind, level, content in _iter_blocks(text):
        if kind == "heading":
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(content)
            run.bold = True
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.size = Pt(18 if level == 1 else 15 if level == 2 else 13)
        else:
            paragraph = document.add_paragraph(content)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _register_pdf_font() -> str:
    font_name = "STSong-Light"
    if font_name in pdfmetrics.getRegisteredFontNames():
        return font_name
    pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    return font_name


def build_pdf(text: str) -> bytes:
    font_name = _register_pdf_font()
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=25 * mm,
        rightMargin=25 * mm,
        topMargin=22 * mm,
        bottomMargin=22 * mm,
        title="文衡优化结果",
    )
    body_style = ParagraphStyle(
        "WenHengBody",
        fontName=font_name,
        fontSize=11,
        leading=19,
        alignment=TA_JUSTIFY,
        wordWrap="CJK",
        firstLineIndent=22,
        spaceAfter=8,
    )
    heading_styles = {
        1: ParagraphStyle("WenHengH1", parent=body_style, fontSize=18, leading=26, alignment=TA_CENTER, firstLineIndent=0, spaceBefore=10, spaceAfter=12),
        2: ParagraphStyle("WenHengH2", parent=body_style, fontSize=15, leading=22, firstLineIndent=0, spaceBefore=10, spaceAfter=8),
        3: ParagraphStyle("WenHengH3", parent=body_style, fontSize=13, leading=20, firstLineIndent=0, spaceBefore=8, spaceAfter=6),
    }
    story = []
    for kind, level, content in _iter_blocks(text):
        escaped = html.escape(content).replace("\n", "<br/>")
        style = heading_styles.get(min(level, 3), body_style) if kind == "heading" else body_style
        story.append(Paragraph(escaped, style))
        story.append(Spacer(1, 2 * mm))
    document.build(story)
    return output.getvalue()
