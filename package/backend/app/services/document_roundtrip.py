import io
import json
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Tuple

import fitz
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.config import get_exe_dir
from app.services.ai_service import split_text_into_segments
from app.word_formatter.utils.ooxml import DocxPackage


SUPPORTED_SOURCE_FORMATS = {"txt", "md", "docx", "pdf"}
UTF8_BOM = b"\xef\xbb\xbf"
TEXT_BOMS = (
    (b"\x00\x00\xfe\xff", "utf-32-be"),
    (b"\xff\xfe\x00\x00", "utf-32-le"),
    (b"\xef\xbb\xbf", "utf-8"),
    (b"\xfe\xff", "utf-16-be"),
    (b"\xff\xfe", "utf-16-le"),
)


def source_document_path(session_id: str, source_format: str) -> str:
    directory = os.path.join(get_exe_dir(), "source_documents")
    os.makedirs(directory, exist_ok=True)
    return os.path.join(directory, f"{session_id}.{source_format}")


def _build_segment_manifest(
    blocks: List[Dict[str, Any]],
) -> Tuple[List[str], List[int], List[str], List[str]]:
    segments: List[str] = []
    segment_blocks: List[int] = []
    segment_prefixes: List[str] = []
    block_suffixes: List[str] = []
    for block_index, block in enumerate(blocks):
        block_text = block["text"]
        cursor = 0
        for segment in split_text_into_segments(block_text):
            position = block_text.find(segment, cursor)
            if position < 0:
                position = cursor
            segment_prefixes.append(block_text[cursor:position])
            segments.append(segment)
            segment_blocks.append(block_index)
            cursor = position + len(segment)
        block_suffixes.append(block_text[cursor:])
    return segments, segment_blocks, segment_prefixes, block_suffixes


def _split_edge_spacing(text: str) -> Tuple[str, str, str]:
    prefix = text[:len(text) - len(text.lstrip())]
    remainder = text[len(prefix):]
    suffix = remainder[len(remainder.rstrip()):]
    body = remainder[:len(remainder) - len(suffix)] if suffix else remainder
    return prefix, body, suffix


def _is_plain_docx_text(paragraph: Paragraph) -> bool:
    return not paragraph._p.xpath(".//w:fldChar | .//w:instrText")


def _decode_source_text(content: bytes) -> Tuple[str, str, bytes]:
    for bom, encoding in TEXT_BOMS:
        if content.startswith(bom):
            return content[len(bom):].decode(encoding), encoding, bom
    for encoding in ("utf-8", "gb18030"):
        try:
            return content.decode(encoding), encoding, b""
        except UnicodeDecodeError:
            continue
    raise ValueError("文本文件编码无法识别，仅支持 UTF-8、UTF-16、UTF-32 或 GB18030")


def parse_text_document(content: bytes, source_format: str) -> Tuple[str, Dict[str, Any], List[str]]:
    text, encoding, bom = _decode_source_text(content)
    blocks: List[Dict[str, Any]] = []
    in_code_fence = False
    for line_index, raw_line in enumerate(text.splitlines(keepends=True)):
        line = raw_line.rstrip("\r\n")
        stripped = line.strip()
        if source_format == "md" and re.match(r"^\s*(```|~~~)", line):
            in_code_fence = not in_code_fence
            continue
        if not stripped or in_code_fence:
            continue
        if source_format == "md" and (line.startswith("    ") or line.startswith("\t")):
            continue

        marker = ""
        body_with_spacing = line
        if source_format == "md":
            match = re.match(r"^(\s*(?:#{1,6}\s+|[-*+]\s+|\d+[.)]\s+|>\s+)?)(.*)$", line)
            if match:
                marker, body_with_spacing = match.groups()
        else:
            marker = line[:len(line) - len(line.lstrip())]
            body_with_spacing = line[len(marker):]

        trailing = body_with_spacing[len(body_with_spacing.rstrip()):]
        body = body_with_spacing[:len(body_with_spacing) - len(trailing)] if trailing else body_with_spacing
        if not body.strip():
            continue
        blocks.append({
            "text": body,
            "locator": {"line": line_index},
            "prefix": marker,
            "suffix": trailing,
        })
    segments, segment_blocks, segment_prefixes, block_suffixes = _build_segment_manifest(blocks)
    return text, {
        "format": source_format,
        "blocks": blocks,
        "segment_blocks": segment_blocks,
        "segment_prefixes": segment_prefixes,
        "block_suffixes": block_suffixes,
        "encoding": encoding,
        "bom_hex": bom.hex(),
    }, segments


def parse_docx_document(content: bytes) -> Tuple[str, Dict[str, Any], List[str]]:
    DocxPackage.from_bytes(content)
    document = Document(io.BytesIO(content))
    blocks: List[Dict[str, Any]] = []

    paragraph_indexes = {
        paragraph._p: paragraph_index
        for paragraph_index, paragraph in enumerate(document.paragraphs)
    }
    table_indexes = {
        table._tbl: table_index
        for table_index, table in enumerate(document.tables)
    }

    for item in document.iter_inner_content():
        if isinstance(item, Paragraph):
            if not _is_plain_docx_text(item):
                continue
            prefix, text, suffix = _split_edge_spacing(item.text)
            if text:
                blocks.append({
                    "text": text,
                    "prefix": prefix,
                    "suffix": suffix,
                    "locator": {
                        "kind": "body",
                        "paragraph": paragraph_indexes[item._p],
                    },
                })
            continue

        if not isinstance(item, Table):
            continue

        table_index = table_indexes[item._tbl]
        seen_cells = set()
        for row_index, row in enumerate(item.rows):
            for cell_index, cell in enumerate(row.cells):
                # python-docx returns the same underlying cell once for every
                # grid position covered by a merge. Process that cell only at
                # its first visual position or its text is duplicated and the
                # export is overwritten repeatedly.
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)

                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    if not _is_plain_docx_text(paragraph):
                        continue
                    prefix, text, suffix = _split_edge_spacing(paragraph.text)
                    if text:
                        blocks.append({
                            "text": text,
                            "prefix": prefix,
                            "suffix": suffix,
                            "locator": {
                                "kind": "table",
                                "table": table_index,
                                "row": row_index,
                                "cell": cell_index,
                                "paragraph": paragraph_index,
                            },
                        })

    segments, segment_blocks, segment_prefixes, block_suffixes = _build_segment_manifest(blocks)
    text = "\n\n".join(block["text"] for block in blocks)
    return text, {
        "format": "docx",
        "blocks": blocks,
        "segment_blocks": segment_blocks,
        "segment_prefixes": segment_prefixes,
        "block_suffixes": block_suffixes,
    }, segments


def parse_pdf_document(content: bytes) -> Tuple[str, Dict[str, Any], List[str]]:
    document = fitz.open(stream=content, filetype="pdf")
    if document.page_count > 100:
        raise ValueError("PDF 页数不能超过 100 页")

    blocks: List[Dict[str, Any]] = []
    for page_index, page in enumerate(document):
        page_data = page.get_text("dict")
        for block in page_data.get("blocks", []):
            if block.get("type") != 0:
                continue
            lines = block.get("lines", [])
            spans = [span for line in lines for span in line.get("spans", []) if span.get("text", "").strip()]
            text = "\n".join(
                "".join(span.get("text", "") for span in line.get("spans", [])).strip()
                for line in lines
            ).strip()
            if not text or not spans:
                continue
            first_span = spans[0]
            line_count = max(len(lines), 1)
            box_height = max(float(block["bbox"][3]) - float(block["bbox"][1]), 1)
            font_size = float(first_span.get("size", 11))
            blocks.append({
                "text": text,
                "locator": {
                    "page": page_index,
                    "bbox": [float(value) for value in block["bbox"]],
                    "font_size": font_size,
                    "font_name": str(first_span.get("font", "")),
                    "color": int(first_span.get("color", 0)),
                    "lineheight": min(max(box_height / max(font_size * line_count, 1), 0.8), 2.0),
                },
            })
    document.close()
    segments, segment_blocks, segment_prefixes, block_suffixes = _build_segment_manifest(blocks)
    text = "\n\n".join(block["text"] for block in blocks)
    return text, {
        "format": "pdf",
        "blocks": blocks,
        "segment_blocks": segment_blocks,
        "segment_prefixes": segment_prefixes,
        "block_suffixes": block_suffixes,
    }, segments


def parse_source_document(content: bytes, source_format: str):
    if source_format in {"txt", "md"}:
        return parse_text_document(content, source_format)
    if source_format == "docx":
        return parse_docx_document(content)
    if source_format == "pdf":
        return parse_pdf_document(content)
    raise ValueError("不支持的文件格式")


def _build_segment_layout(
    manifest: Dict[str, Any],
    original_segments: List[str] | None,
) -> Tuple[List[str], List[str]]:
    prefixes = manifest.get("segment_prefixes")
    suffixes = manifest.get("block_suffixes")
    segment_blocks = manifest.get("segment_blocks", [])
    block_count = len(manifest.get("blocks", []))
    if (
        isinstance(prefixes, list)
        and len(prefixes) == len(segment_blocks)
        and isinstance(suffixes, list)
        and len(suffixes) == block_count
    ):
        return prefixes, suffixes

    inferred_prefixes = ["" for _ in segment_blocks]
    inferred_suffixes = ["" for _ in range(block_count)]
    if not original_segments:
        return inferred_prefixes, inferred_suffixes

    by_block: List[List[int]] = [[] for _ in range(block_count)]
    for segment_index, block_index in enumerate(segment_blocks):
        if 0 <= block_index < block_count and segment_index < len(original_segments):
            by_block[block_index].append(segment_index)
    for block_index, segment_indexes in enumerate(by_block):
        block_text = manifest["blocks"][block_index]["text"]
        cursor = 0
        for segment_index in segment_indexes:
            segment = original_segments[segment_index]
            position = block_text.find(segment, cursor)
            if position < 0:
                position = cursor
            inferred_prefixes[segment_index] = block_text[cursor:position]
            cursor = position + len(segment)
        inferred_suffixes[block_index] = block_text[cursor:]
    return inferred_prefixes, inferred_suffixes


def _group_optimized_blocks(
    manifest: Dict[str, Any],
    optimized_segments: List[str],
    original_segments: List[str] | None = None,
) -> List[str]:
    block_count = len(manifest.get("blocks", []))
    grouped: List[List[str]] = [[] for _ in range(block_count)]
    prefixes, suffixes = _build_segment_layout(manifest, original_segments)
    for segment_index, block_index in enumerate(manifest.get("segment_blocks", [])):
        if segment_index < len(optimized_segments) and 0 <= block_index < block_count:
            grouped[block_index].append(
                f"{prefixes[segment_index]}{optimized_segments[segment_index]}"
            )
    return [
        f"{''.join(parts)}{suffixes[index]}" if parts else manifest["blocks"][index]["text"]
        for index, parts in enumerate(grouped)
    ]


def _get_docx_paragraph(document: Document, locator: Dict[str, Any]):
    if locator["kind"] == "body":
        return document.paragraphs[locator["paragraph"]]
    return document.tables[locator["table"]].rows[locator["row"]].cells[locator["cell"]].paragraphs[locator["paragraph"]]


def _replace_runs_preserving_style(paragraph, text: str):
    text_nodes = list(paragraph._p.xpath(".//w:t"))
    if not text_nodes:
        paragraph.add_run(text)
        return
    original_lengths = [max(len(node.text or ""), 1) for node in text_nodes]
    total = sum(original_lengths)
    cursor = 0
    consumed_weight = 0
    for index, node in enumerate(text_nodes):
        consumed_weight += original_lengths[index]
        end = len(text) if index == len(text_nodes) - 1 else round(len(text) * consumed_weight / total)
        replacement = text[cursor:end]
        node.text = replacement
        if replacement.startswith(" ") or replacement.endswith(" "):
            node.set(qn("xml:space"), "preserve")
        else:
            node.attrib.pop(qn("xml:space"), None)
        cursor = end


def build_text_from_source(
    source_path: str,
    manifest_json: str,
    optimized_segments: List[str],
    original_segments: List[str] | None = None,
) -> bytes:
    manifest = json.loads(manifest_json)
    source = Path(source_path).read_bytes()
    encoding = manifest.get("encoding", "utf-8")
    bom_hex = manifest.get("bom_hex")
    if bom_hex is None:
        bom = UTF8_BOM if manifest.get("utf8_bom") else b""
    else:
        try:
            bom = bytes.fromhex(bom_hex)
        except ValueError as exc:
            raise ValueError("源文本编码清单无效") from exc
    text = source[len(bom):].decode(encoding)
    lines = text.splitlines(keepends=True)
    optimized_blocks = _group_optimized_blocks(manifest, optimized_segments, original_segments)
    for block, optimized_text in zip(manifest.get("blocks", []), optimized_blocks):
        line_index = block.get("locator", {}).get("line")
        if line_index is None or not 0 <= line_index < len(lines):
            continue
        raw_line = lines[line_index]
        ending_match = re.search(r"(\r\n|\n|\r)$", raw_line)
        ending = ending_match.group(1) if ending_match else ""
        lines[line_index] = f"{block.get('prefix', '')}{optimized_text}{block.get('suffix', '')}{ending}"
    return bom + "".join(lines).encode(encoding)


def build_docx_from_source(
    source_path: str,
    manifest_json: str,
    optimized_segments: List[str],
    original_segments: List[str] | None = None,
) -> bytes:
    manifest = json.loads(manifest_json)
    document = Document(source_path)
    optimized_blocks = _group_optimized_blocks(manifest, optimized_segments, original_segments)
    for block, text in zip(manifest["blocks"], optimized_blocks):
        paragraph = _get_docx_paragraph(document, block["locator"])
        _replace_runs_preserving_style(
            paragraph,
            f"{block.get('prefix', '')}{text.replace(chr(13), '').replace(chr(10), '').replace(chr(9), '')}{block.get('suffix', '')}",
        )
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _pdf_color(value: int) -> Tuple[float, float, float]:
    return (
        ((value >> 16) & 255) / 255,
        ((value >> 8) & 255) / 255,
        (value & 255) / 255,
    )


def _pdf_replacement_rect(
    locator: Dict[str, Any],
    manifest_blocks: List[Dict[str, Any]],
    page_rect: fitz.Rect,
) -> fitz.Rect:
    rect = fitz.Rect(locator["bbox"])
    safe_bottom = page_rect.y1
    width = max(rect.width, 1)
    for block in manifest_blocks:
        other = block.get("locator", {})
        if other.get("page") != locator.get("page") or other is locator:
            continue
        other_rect = fitz.Rect(other.get("bbox", rect))
        if other_rect.y0 < rect.y1 - 0.5:
            continue
        overlap = max(0, min(rect.x1, other_rect.x1) - max(rect.x0, other_rect.x0))
        if overlap >= min(width, max(other_rect.width, 1)) * 0.25:
            safe_bottom = min(safe_bottom, other_rect.y0 - 1)

    font_size = min(max(locator.get("font_size", 11), 6), 24)
    rect.y1 = max(rect.y1, min(safe_bottom, rect.y1 + font_size * 0.75))
    return rect


def build_pdf_from_source(
    source_path: str,
    manifest_json: str,
    optimized_segments: List[str],
    original_segments: List[str] | None = None,
) -> bytes:
    manifest = json.loads(manifest_json)
    optimized_blocks = _group_optimized_blocks(manifest, optimized_segments, original_segments)
    changed_blocks = [
        (block, text)
        for block, text in zip(manifest["blocks"], optimized_blocks)
        if text != block["text"]
    ]
    if not changed_blocks:
        return Path(source_path).read_bytes()

    document = fitz.open(source_path)
    font_path = next((
        path for path in (
            "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
            "C:/Windows/Fonts/simhei.ttf",
        )
        if os.path.exists(path)
    ), None)
    if not font_path:
        raise RuntimeError("服务器缺少可用于 PDF 回写的中文字体")

    page_blocks: Dict[int, List[Tuple[Dict[str, Any], str]]] = {}
    for block, text in changed_blocks:
        locator = block["locator"]
        page_blocks.setdefault(locator["page"], []).append((locator, text))

    for page_index, items in page_blocks.items():
        page = document[page_index]
        for locator, _ in items:
            page.add_redact_annot(
                fitz.Rect(locator["bbox"]),
                fill=None,
                cross_out=False,
            )
        page.apply_redactions(
            images=fitz.PDF_REDACT_IMAGE_NONE,
            graphics=fitz.PDF_REDACT_LINE_ART_NONE,
        )
        page.insert_font(fontname="WenHengCJK", fontfile=font_path)
        for locator, text in items:
            rect = _pdf_replacement_rect(locator, manifest["blocks"], page.rect)
            font_size = min(max(locator.get("font_size", 11), 6), 24)
            remaining = page.insert_textbox(
                rect,
                text,
                fontname="WenHengCJK",
                fontsize=font_size,
                color=_pdf_color(locator.get("color", 0)),
                lineheight=locator.get("lineheight", 1.15),
            )
            if remaining < 0:
                document.close()
                raise ValueError(
                    f"第 {page_index + 1} 页的润色文字无法在原字号和原文字框内完整放置；"
                    "为避免改变版式，已停止导出。请缩短该段文字后重试。"
                )

    output = document.tobytes(garbage=4, deflate=True)
    document.close()
    return output


def delete_source_document(session_id: str, source_format: str):
    if not source_format:
        return
    path = Path(source_document_path(session_id, source_format))
    if path.exists():
        path.unlink()
