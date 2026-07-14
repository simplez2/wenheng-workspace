import io
import json
import tempfile
import unittest
from pathlib import Path

import fitz
from docx import Document

from app.services.document_roundtrip import (
    build_docx_from_source,
    build_pdf_from_source,
    build_text_from_source,
    parse_docx_document,
    parse_pdf_document,
    parse_text_document,
)


def _fixture_document() -> bytes:
    document = Document()
    document.add_paragraph("正文之前")
    table = document.add_table(rows=2, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "合并单元格中的长段落，需要且只能处理一次。"
    table.cell(1, 0).text = "左下内容"
    table.cell(1, 1).text = "右下内容"
    document.add_paragraph("正文之后")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


class DocumentRoundtripTests(unittest.TestCase):
    def test_txt_roundtrip_preserves_bom_crlf_blank_lines_and_spacing(self):
        content = b"\xef\xbb\xbf" + "  第一行  \r\n\r\n第二行\r\n".encode("utf-8")
        _, manifest, segments = parse_text_document(content, "txt")
        self.assertEqual(segments, ["第一行", "第二行"])

        with tempfile.TemporaryDirectory() as directory:
            source_path = Path(directory) / "source.txt"
            source_path.write_bytes(content)
            result = build_text_from_source(
                str(source_path),
                json.dumps(manifest, ensure_ascii=False),
                ["优化一", "优化二"],
            )

        self.assertEqual(
            result,
            b"\xef\xbb\xbf" + "  优化一  \r\n\r\n优化二\r\n".encode("utf-8"),
        )

    def test_txt_roundtrip_preserves_gb18030_encoding(self):
        content = "第一行\r\n第二行\r\n".encode("gb18030")
        _, manifest, segments = parse_text_document(content, "txt")
        self.assertEqual(manifest["encoding"], "gb18030")
        self.assertEqual(segments, ["第一行", "第二行"])

        with tempfile.TemporaryDirectory() as directory:
            source_path = Path(directory) / "source.txt"
            source_path.write_bytes(content)
            result = build_text_from_source(
                str(source_path),
                json.dumps(manifest, ensure_ascii=False),
                ["优化一", "优化二"],
            )

        self.assertEqual(result, "优化一\r\n优化二\r\n".encode("gb18030"))

    def test_txt_roundtrip_preserves_utf16le_bom(self):
        content = b"\xff\xfe" + "第一行\r\n第二行\r\n".encode("utf-16-le")
        _, manifest, segments = parse_text_document(content, "txt")
        self.assertEqual(manifest["encoding"], "utf-16-le")
        self.assertEqual(manifest["bom_hex"], "fffe")
        self.assertEqual(segments, ["第一行", "第二行"])

        with tempfile.TemporaryDirectory() as directory:
            source_path = Path(directory) / "source.txt"
            source_path.write_bytes(content)
            result = build_text_from_source(
                str(source_path),
                json.dumps(manifest, ensure_ascii=False),
                ["优化一", "优化二"],
            )

        self.assertEqual(result, b"\xff\xfe" + "优化一\r\n优化二\r\n".encode("utf-16-le"))

    def test_markdown_roundtrip_preserves_markers_and_code_blocks(self):
        content = "# 标题\r\n\r\n- 列表项\r\n```python\r\nx = 1\r\n```\r\n".encode("utf-8")
        _, manifest, segments = parse_text_document(content, "md")
        self.assertEqual(segments, ["标题", "列表项"])

        with tempfile.TemporaryDirectory() as directory:
            source_path = Path(directory) / "source.md"
            source_path.write_bytes(content)
            result = build_text_from_source(
                str(source_path),
                json.dumps(manifest, ensure_ascii=False),
                ["新标题", "新列表项"],
            )

        self.assertEqual(
            result.decode("utf-8"),
            "# 新标题\r\n\r\n- 新列表项\r\n```python\r\nx = 1\r\n```\r\n",
        )

    def test_docx_blocks_follow_document_order_and_skip_merged_duplicates(self):
        content = _fixture_document()
        text, manifest, segments = parse_docx_document(content)

        self.assertEqual(
            [block["text"] for block in manifest["blocks"]],
            [
                "正文之前",
                "合并单元格中的长段落，需要且只能处理一次。",
                "左下内容",
                "右下内容",
                "正文之后",
            ],
        )
        self.assertEqual(text.count("合并单元格中的长段落"), 1)
        self.assertEqual(len(segments), 5)

    def test_docx_export_updates_merged_cell_once(self):
        content = _fixture_document()
        _, manifest, segments = parse_docx_document(content)
        optimized = [f"优化：{segment}" for segment in segments]

        with tempfile.TemporaryDirectory() as directory:
            source_path = Path(directory) / "source.docx"
            source_path.write_bytes(content)
            result = build_docx_from_source(
                str(source_path),
                json.dumps(manifest, ensure_ascii=False),
                optimized,
            )

        document = Document(io.BytesIO(result))
        self.assertEqual(document.paragraphs[0].text, "优化：正文之前")
        self.assertEqual(
            document.tables[0].cell(0, 0).text,
            "优化：合并单元格中的长段落，需要且只能处理一次。",
        )
        self.assertEqual(document.paragraphs[-1].text, "优化：正文之后")

    def test_docx_export_preserves_run_and_paragraph_formatting(self):
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_with_next = True
        first = paragraph.add_run("前半部分")
        first.bold = True
        second = paragraph.add_run("后半部分")
        second.italic = True
        source = io.BytesIO()
        document.save(source)

        _, manifest, segments = parse_docx_document(source.getvalue())
        self.assertEqual(len(segments), 1)
        with tempfile.TemporaryDirectory() as directory:
            source_path = Path(directory) / "source.docx"
            source_path.write_bytes(source.getvalue())
            result = build_docx_from_source(
                str(source_path),
                json.dumps(manifest, ensure_ascii=False),
                ["润色后仍然保留原来的运行样式"],
            )

        exported = Document(io.BytesIO(result))
        exported_paragraph = exported.paragraphs[0]
        self.assertEqual(exported_paragraph.text, "润色后仍然保留原来的运行样式")
        self.assertTrue(exported_paragraph.paragraph_format.keep_with_next)
        self.assertEqual(len(exported_paragraph.runs), 2)
        self.assertTrue(exported_paragraph.runs[0].bold)
        self.assertTrue(exported_paragraph.runs[1].italic)

    def test_docx_export_preserves_manual_line_breaks(self):
        document = Document()
        run = document.add_paragraph().add_run("第一行")
        run.add_break()
        run.add_text("第二行")
        source = io.BytesIO()
        document.save(source)

        _, manifest, segments = parse_docx_document(source.getvalue())
        self.assertEqual(segments, ["第一行", "第二行"])
        with tempfile.TemporaryDirectory() as directory:
            source_path = Path(directory) / "source.docx"
            source_path.write_bytes(source.getvalue())
            result = build_docx_from_source(
                str(source_path),
                json.dumps(manifest, ensure_ascii=False),
                ["优化一", "优化二"],
            )

        exported = Document(io.BytesIO(result))
        self.assertEqual(exported.paragraphs[0].text, "优化一\n优化二")
        self.assertEqual(len(exported.paragraphs[0]._p.xpath(".//w:br")), 1)

    def test_pdf_identity_export_returns_original_bytes(self):
        document = fitz.open()
        page = document.new_page(width=300, height=200)
        page.insert_textbox(
            fitz.Rect(40, 40, 260, 100),
            "Original PDF text\nSecond line",
            fontsize=11,
        )
        content = document.tobytes()
        document.close()

        _, manifest, segments = parse_pdf_document(content)
        with tempfile.TemporaryDirectory() as directory:
            source_path = Path(directory) / "source.pdf"
            source_path.write_bytes(content)
            result = build_pdf_from_source(
                str(source_path),
                json.dumps(manifest, ensure_ascii=False),
                segments,
            )

        self.assertEqual(result, content)


if __name__ == "__main__":
    unittest.main()
