import base64
import io
import unittest
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt
from lxml import etree

from app.word_formatter.services.compiler import CompileOptions
from app.word_formatter.services.existing_docx_formatter import format_existing_docx


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP = {"w": W_NS}
PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wl2nWQAAAAASUVORK5CYII="
)


def _document_xml(content: bytes):
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        return etree.fromstring(archive.read("word/document.xml"))


def _signature(content: bytes):
    root = _document_xml(content)
    return {
        "tables": len(root.findall(".//w:tbl", namespaces=NSMAP)),
        "sections": len(root.findall(".//w:sectPr", namespaces=NSMAP)),
        "drawings": len(root.findall(".//w:drawing", namespaces=NSMAP)),
        "paragraphs": len(root.findall(".//w:p", namespaces=NSMAP)),
        "text": "".join(root.xpath(".//w:t/text()", namespaces=NSMAP)),
        "grid_spans": [
            item.get(f"{{{W_NS}}}val")
            for item in root.findall(".//w:gridSpan", namespaces=NSMAP)
        ],
    }


def _table_run_properties(content: bytes):
    root = _document_xml(content)
    properties = []
    for run in root.findall(".//w:tbl//w:r", namespaces=NSMAP):
        text = "".join(run.xpath(".//w:t/text()", namespaces=NSMAP))
        if not text:
            continue
        r_pr = run.find("w:rPr", namespaces=NSMAP)

        def value(local, attribute="val"):
            if r_pr is None:
                return None
            element = r_pr.find(f"w:{local}", namespaces=NSMAP)
            return None if element is None else element.get(f"{{{W_NS}}}{attribute}")

        fonts = None if r_pr is None else r_pr.find("w:rFonts", namespaces=NSMAP)
        properties.append(
            {
                "text": text,
                "size": value("sz"),
                "size_cs": value("szCs"),
                "bold": value("b"),
                "ascii": None if fonts is None else fonts.get(f"{{{W_NS}}}ascii"),
                "hansi": None if fonts is None else fonts.get(f"{{{W_NS}}}hAnsi"),
                "east_asia": None if fonts is None else fonts.get(f"{{{W_NS}}}eastAsia"),
            }
        )
    return properties


def _structured_fixture() -> bytes:
    document = Document()
    document.add_paragraph("2026 template cover")
    table = document.add_table(rows=6, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "Merged template heading"
    for index in range(1, 6):
        table.cell(index, 0).text = f"Label {index}"
        table.cell(index, 1).text = (
            "This is a long narrative paragraph inside the original template cell. "
            "It must be formatted without rebuilding the table or duplicating text."
        )
    picture_paragraph = document.add_paragraph()
    picture_paragraph.add_run().add_picture(io.BytesIO(PNG_1X1))
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Second section text remains in the same section structure.")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _structured_table_font_fixture() -> bytes:
    document = Document()
    document.add_paragraph("Template document")
    table = document.add_table(rows=6, cols=2)
    texts = [
        "\u73af\u8282",
        "\u4e3b\u8981\u6d3b\u52a8",
        "\u8bfb\u5883\u8bc6\u9898",
        "\u9605\u8bfb\u8fd0\u52a8\u6545\u4e8b\u3001\u89c4\u5219\u3001\u5065\u5eb7\u8bf4\u660e\u3001\u4f53\u8d28\u6570\u636e\u548c\u56fe\u8868",
        "\u7acb\u6807\u89c4\u5212",
        "\u9009\u62e9\u9002\u5207\u76ee\u6807\u3001\u89d2\u8272\u548c\u4efb\u52a1\u4ea7\u54c1\uff0c\u5236\u5b9a\u884c\u52a8\u8ba1\u5212",
        "\u4eb2\u5386\u8fd0\u52a8",
        "\u5728\u5b89\u5168\u3001\u8db3\u91cf\u7684\u8eab\u4f53\u7ec3\u4e60\u4e2d\u4f53\u9a8c\u52a8\u4f5c\u3001\u89c4\u5219\u4e0e\u5408\u4f5c",
        "\u53d6\u8bc1\u8bb0\u5f55",
        "\u8bb0\u5f55\u7ec3\u4e60\u6b21\u6570\u3001\u611f\u53d7\u3001\u9519\u8bef\u7c7b\u578b\u3001\u5fc3\u7387\u6216\u4f53\u6d4b\u53d8\u5316\u7b49",
        "\u8868\u8fbe\u4ea4\u6d41",
        "\u53e3\u5934\u8bf4\u660e\u3001\u52a8\u4f5c\u89e3\u8bf4\u3001\u56fe\u6587\u8bf4\u660e\u3001\u91c7\u8bbf\u3001\u65b0\u95fb\u3001\u62a5\u544a\u6216\u89c6\u9891",
    ]
    for cell, text in zip((cell for row in table.rows for cell in row.cells), texts):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run._element.get_or_add_rPr().rFonts.set(f"{{{W_NS}}}eastAsia", "SimSun")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


class ExistingDocxFormatterTests(unittest.TestCase):
    def test_structured_docx_preserves_content_and_layout_objects(self):
        source = _structured_fixture()
        before = _signature(source)

        result = format_existing_docx(
            source,
            CompileOptions(include_cover=True, include_toc=True),
        )

        self.assertTrue(result.success, result.error)
        after = _signature(result.docx_bytes)
        self.assertEqual(after, before)
        self.assertTrue(any("模板型 Word" in item for item in result.warnings))

    def test_output_is_openable_by_python_docx(self):
        result = format_existing_docx(_structured_fixture(), CompileOptions())
        self.assertTrue(result.success, result.error)
        document = Document(io.BytesIO(result.docx_bytes))
        self.assertEqual(len(document.tables), 1)
        self.assertEqual(len(document.sections), 2)
        self.assertEqual(document.tables[0].cell(0, 0).text, "Merged template heading")

    def test_structured_docx_preserves_table_run_typography(self):
        source = _structured_table_font_fixture()
        before = _table_run_properties(source)

        result = format_existing_docx(source, CompileOptions())

        self.assertTrue(result.success, result.error)
        self.assertEqual(_table_run_properties(result.docx_bytes), before)


if __name__ == "__main__":
    unittest.main()
